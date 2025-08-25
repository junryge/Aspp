# -*- coding: utf-8 -*-
"""
Korean Tech Document RAG System with GGUF - Complete Offline Version
폐쇄망 환경을 위한 완전한 오프라인 한국어 RAG 시스템
"""

import os
import sys
import json
import time
import threading
import queue
import hashlib
import re
import pickle
import numpy as np
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List, Any, Generator, Tuple
import logging
from dataclasses import dataclass, field
from enum import Enum
from collections import Counter, defaultdict
import math

# UI 라이브러리
import customtkinter as ctk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import tkinter as tk

# GGUF 모델 라이브러리
try:
    from llama_cpp import Llama
    LLAMA_CPP_AVAILABLE = True
except ImportError:
    LLAMA_CPP_AVAILABLE = False
    print("⚠️ llama-cpp-python이 설치되지 않았습니다.")
    print("설치: pip install llama-cpp-python")

# 문서 처리 라이브러리
try:
    from docx import Document
    import faiss
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    RAG_LIBS_AVAILABLE = True
except ImportError:
    RAG_LIBS_AVAILABLE = False
    print("⚠️ RAG 라이브러리가 설치되지 않았습니다.")
    print("설치: pip install python-docx faiss-cpu scikit-learn")

# 로깅 설정
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('gguf_rag_offline.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# RAG 프롬프트 템플릿
class RAGPromptTemplate(Enum):
    """한글 기술문서 RAG 전용 프롬프트 템플릿"""
    
    KOREAN_TECH_RAG = """당신은 기술 문서를 정확하게 분석하고 답변하는 AI 전문가입니다.

[시스템 지침]
{system_prompt}

[참조 문서]
{context}

[중요 지침]
1. 위 참조 문서의 내용만을 기반으로 답변하세요.
2. 문서에 없는 내용은 추측하지 마세요.
3. SQL 쿼리나 코드가 있다면 정확히 인용하세요.
4. 기술 용어는 원문 그대로 사용하세요.
5. 답변의 출처를 명시하세요.

[대화 기록]
{chat_history}

사용자: {user_message}
AI 전문가:"""

    QUERY_FOCUSED_RAG = """당신은 SQL 및 기술 쿼리 전문가입니다.

[전문 분야]
- SQL 쿼리 작성 및 최적화
- 데이터베이스 설계
- 기술 문서 해석

[참조 문서 및 예제]
{context}

[쿼리 작성 규칙]
1. 문서의 스키마와 규칙을 정확히 따르세요.
2. 한글 컬럼명이 있다면 그대로 사용하세요.
3. 성능을 고려한 쿼리를 작성하세요.
4. 주석으로 설명을 추가하세요.

{chat_history}

요청사항: {user_message}
응답:"""

@dataclass
class DocumentChunk:
    """문서 청크 데이터 클래스"""
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[np.ndarray] = None
    chunk_id: str = ""
    
    def __post_init__(self):
        if not self.chunk_id:
            self.chunk_id = hashlib.md5(self.content.encode()).hexdigest()[:8]

@dataclass
class RAGConfig:
    """RAG 설정 데이터 클래스"""
    chunk_size: int = 1000
    chunk_overlap: int = 200
    max_chunks_per_query: int = 5
    embedding_model_name: str = "offline"  # 오프라인 모드
    use_hybrid_search: bool = True
    min_relevance_score: float = 0.3
    embedding_dimension: int = 768

class KoreanOfflineEmbedding:
    """폐쇄망용 오프라인 한국어 임베딩 시스템"""
    
    def __init__(self, dimension: int = 768):
        self.dimension = dimension
        self.vocab = {}
        self.idf_scores = {}
        self.word_vectors = None
        self.char_vocab = {}
        self.subword_vocab = {}
        self.document_count = 0
        self.trained = False
        
    def build_from_documents(self, documents: List[str]):
        """문서 집합으로부터 임베딩 모델 구축"""
        logger.info("오프라인 임베딩 모델 학습 시작...")
        
        self.document_count = len(documents)
        
        # 1. 어휘 구축
        self._build_vocabulary(documents)
        
        # 2. IDF 계산
        self._calculate_idf(documents)
        
        # 3. 단어 벡터 초기화
        self._initialize_word_vectors()
        
        # 4. 간단한 Word2Vec 학습 (선택적)
        if len(documents) > 10:
            self._train_simple_word2vec(documents)
        
        self.trained = True
        logger.info(f"✓ 임베딩 모델 학습 완료! 어휘 크기: {len(self.vocab)}")
        
    def _build_vocabulary(self, documents: List[str]):
        """어휘 구축"""
        word_freq = Counter()
        char_freq = Counter()
        
        for doc in documents:
            tokens = self._tokenize_korean(doc)
            word_freq.update(tokens)
            
            for token in tokens:
                char_freq.update(token)
        
        # 빈도 2 이상인 단어만 사용
        self.vocab = {word: idx for idx, (word, freq) in 
                     enumerate(word_freq.most_common(50000)) if freq >= 2}
        
        # 문자 어휘
        self.char_vocab = {char: idx for idx, (char, _) in 
                          enumerate(char_freq.most_common(5000))}
        
        logger.info(f"어휘 크기: {len(self.vocab)}, 문자 어휘: {len(self.char_vocab)}")
        
    def _tokenize_korean(self, text: str) -> List[str]:
        """한국어 토크나이저 (형태소 분석기 없이)"""
        text = text.lower().strip()
        
        # 기본 토큰화
        tokens = re.findall(r'[가-힣]+|[a-zA-Z]+|[0-9]+', text)
        
        # 간단한 조사 분리
        korean_tokens = []
        josa_patterns = ['은', '는', '이', '가', '을', '를', '에', '에서', 
                        '으로', '와', '과', '의', '도', '만', '부터', '까지']
        
        for token in tokens:
            if re.match(r'^[가-힣]+$', token) and len(token) > 2:
                for josa in josa_patterns:
                    if token.endswith(josa) and len(token) > len(josa) + 1:
                        word = token[:-len(josa)]
                        korean_tokens.extend([word, josa])
                        break
                else:
                    korean_tokens.append(token)
            else:
                korean_tokens.append(token)
                
        return korean_tokens
        
    def _calculate_idf(self, documents: List[str]):
        """IDF 계산"""
        doc_count = len(documents)
        word_doc_freq = defaultdict(int)
        
        for doc in documents:
            tokens = set(self._tokenize_korean(doc))
            for token in tokens:
                if token in self.vocab:
                    word_doc_freq[token] += 1
        
        for word, doc_freq in word_doc_freq.items():
            self.idf_scores[word] = math.log(doc_count / (doc_freq + 1))
            
    def _initialize_word_vectors(self):
        """단어 벡터 초기화"""
        vocab_size = len(self.vocab)
        
        # Xavier 초기화
        scale = np.sqrt(2.0 / self.dimension)
        self.word_vectors = np.random.normal(0, scale, (vocab_size, self.dimension))
        
        # 문자 기반 초기화 보강
        for word, idx in self.vocab.items():
            char_vector = self._get_char_vector(word)
            self.word_vectors[idx] = 0.7 * self.word_vectors[idx] + 0.3 * char_vector
            
    def _get_char_vector(self, word: str) -> np.ndarray:
        """문자 기반 벡터 생성"""
        vector = np.zeros(self.dimension)
        
        for i, char in enumerate(word):
            if char in self.char_vocab:
                # 문자와 위치 정보를 조합한 해싱
                seed = f"{char}_{i}_{word}"
                hash_val = int(hashlib.md5(seed.encode()).hexdigest(), 16)
                
                # 여러 위치에 분산
                for j in range(3):
                    idx = (hash_val + j * 1000) % self.dimension
                    vector[idx] += 1.0 / (i + 1)  # 위치에 따른 가중치
        
        # 정규화
        norm = np.linalg.norm(vector)
        if norm > 0:
            vector /= norm
            
        return vector
        
    def _train_simple_word2vec(self, documents: List[str], epochs: int = 2):
        """간단한 Word2Vec 스타일 학습"""
        logger.info("Word2Vec 학습 중...")
        
        learning_rate = 0.025
        window_size = 5
        
        for epoch in range(epochs):
            logger.info(f"  Epoch {epoch + 1}/{epochs}")
            
            for doc_idx, doc in enumerate(documents):
                if doc_idx % 100 == 0:
                    logger.info(f"    문서 {doc_idx}/{len(documents)}")
                    
                tokens = self._tokenize_korean(doc)
                token_indices = [self.vocab.get(token, -1) for token in tokens]
                token_indices = [idx for idx in token_indices if idx != -1]
                
                for i, center_idx in enumerate(token_indices):
                    context_indices = []
                    
                    for j in range(max(0, i - window_size), min(len(token_indices), i + window_size + 1)):
                        if i != j:
                            context_indices.append(token_indices[j])
                    
                    if context_indices:
                        # 간단한 업데이트
                        center_vec = self.word_vectors[center_idx]
                        
                        for context_idx in context_indices:
                            context_vec = self.word_vectors[context_idx]
                            
                            # 코사인 유사도
                            similarity = np.dot(center_vec, context_vec)
                            
                            # 그래디언트 업데이트
                            gradient = learning_rate * (1 - similarity) * context_vec
                            self.word_vectors[center_idx] += gradient
                            self.word_vectors[context_idx] += gradient * 0.5
            
            # 학습률 감소
            learning_rate *= 0.9
            
            # 정규화
            norms = np.linalg.norm(self.word_vectors, axis=1, keepdims=True)
            self.word_vectors /= (norms + 1e-8)
        
        logger.info("✓ Word2Vec 학습 완료")
        
    def encode(self, texts: Union[str, List[str]], show_progress_bar: bool = False) -> np.ndarray:
        """텍스트를 임베딩으로 변환"""
        if isinstance(texts, str):
            texts = [texts]
            
        embeddings = []
        
        for text in texts:
            # 1. TF-IDF 벡터
            tfidf_vector = self._get_tfidf_vector(text)
            
            # 2. 평균 단어 벡터
            avg_word_vector = self._get_avg_word_vector(text)
            
            # 3. 문자 기반 벡터
            char_vector = self._get_char_vector(text[:100])  # 처음 100자만
            
            # 결합 (가중 평균)
            combined_vector = (
                0.4 * tfidf_vector + 
                0.4 * avg_word_vector + 
                0.2 * char_vector
            )
            
            # 최종 정규화
            norm = np.linalg.norm(combined_vector)
            if norm > 0:
                combined_vector /= norm
                
            embeddings.append(combined_vector)
            
        return np.array(embeddings)
        
    def _get_tfidf_vector(self, text: str) -> np.ndarray:
        """TF-IDF 벡터 생성"""
        vector = np.zeros(self.dimension)
        tokens = self._tokenize_korean(text)
        
        if not tokens:
            return vector
            
        # TF 계산
        token_counts = Counter(tokens)
        total_tokens = len(tokens)
        
        # TF-IDF 벡터 생성
        for token, count in token_counts.items():
            if token in self.vocab:
                tf = count / total_tokens
                idf = self.idf_scores.get(token, 0)
                tfidf = tf * idf
                
                # 여러 차원에 분산
                base_idx = self.vocab[token]
                indices = [
                    base_idx % self.dimension,
                    (base_idx * 7) % self.dimension,
                    (base_idx * 13) % self.dimension
                ]
                
                for idx in indices:
                    vector[idx] += tfidf
                    
        return vector
        
    def _get_avg_word_vector(self, text: str) -> np.ndarray:
        """평균 단어 벡터 생성"""
        if self.word_vectors is None:
            return np.zeros(self.dimension)
            
        tokens = self._tokenize_korean(text)
        vectors = []
        
        for token in tokens:
            if token in self.vocab:
                idx = self.vocab[token]
                vectors.append(self.word_vectors[idx])
                
        if vectors:
            return np.mean(vectors, axis=0)
        else:
            return np.zeros(self.dimension)
            
    def save_model(self, path: str):
        """모델 저장"""
        model_data = {
            'vocab': self.vocab,
            'idf_scores': self.idf_scores,
            'word_vectors': self.word_vectors,
            'char_vocab': self.char_vocab,
            'dimension': self.dimension,
            'document_count': self.document_count
        }
        
        with open(path, 'wb') as f:
            pickle.dump(model_data, f)
            
        logger.info(f"모델 저장 완료: {path}")
        
    def load_model(self, path: str):
        """모델 로드"""
        with open(path, 'rb') as f:
            model_data = pickle.load(f)
            
        self.vocab = model_data['vocab']
        self.idf_scores = model_data['idf_scores']
        self.word_vectors = model_data['word_vectors']
        self.char_vocab = model_data['char_vocab']
        self.dimension = model_data['dimension']
        self.document_count = model_data['document_count']
        self.trained = True
        
        logger.info(f"모델 로드 완료: {path}")

class KoreanTextSplitter:
    """한글 문서 특화 텍스트 분할기"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
    def split_text(self, text: str) -> List[str]:
        """한글 문서를 의미 단위로 분할"""
        # 문단 우선 분할
        paragraphs = text.split('\n\n')
        
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            # SQL 쿼리나 코드 블록은 분할하지 않음
            if self._is_code_block(para):
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                chunks.append(para.strip())
                continue
                
            # 일반 텍스트 처리
            sentences = self._split_korean_sentences(para)
            
            for sentence in sentences:
                if len(current_chunk) + len(sentence) < self.chunk_size:
                    current_chunk += sentence + " "
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence + " "
                    
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        # 오버랩 적용
        return self._apply_overlap(chunks)
        
    def _split_korean_sentences(self, text: str) -> List[str]:
        """한글 문장 분할"""
        # 한글 문장 종결 패턴
        sentence_enders = r'[.!?。！？][\s"]'
        sentences = re.split(sentence_enders, text)
        
        # 빈 문장 제거 및 정리
        return [s.strip() for s in sentences if s.strip()]
        
    def _is_code_block(self, text: str) -> bool:
        """코드 블록 감지"""
        code_indicators = [
            'SELECT', 'FROM', 'WHERE', 'CREATE TABLE',
            'INSERT INTO', 'UPDATE', 'DELETE',
            '```', 'def ', 'class ', 'function'
        ]
        return any(indicator in text.upper() for indicator in code_indicators)
        
    def _apply_overlap(self, chunks: List[str]) -> List[str]:
        """청크 간 오버랩 적용"""
        if not chunks or len(chunks) <= 1:
            return chunks
            
        overlapped_chunks = []
        for i, chunk in enumerate(chunks):
            if i > 0:
                # 이전 청크의 마지막 부분 추가
                prev_chunk = chunks[i-1]
                overlap_text = prev_chunk[-self.chunk_overlap:] if len(prev_chunk) > self.chunk_overlap else prev_chunk
                chunk = overlap_text + " " + chunk
            overlapped_chunks.append(chunk)
            
        return overlapped_chunks

class DocumentProcessor:
    """문서 처리 및 임베딩 관리"""
    
    def __init__(self, rag_config: RAGConfig):
        self.config = rag_config
        self.text_splitter = KoreanTextSplitter(
            chunk_size=rag_config.chunk_size,
            chunk_overlap=rag_config.chunk_overlap
        )
        self.embedding_model = KoreanOfflineEmbedding(rag_config.embedding_dimension)
        self.chunks_db = []
        self.embeddings = None
        self.index = None
        self.tfidf_vectorizer = None
        self.tfidf_matrix = None
        
    def initialize_embedding_model(self, model_path: Optional[str] = None) -> bool:
        """임베딩 모델 초기화"""
        try:
            if model_path and os.path.exists(model_path):
                # 저장된 모델 로드
                self.embedding_model.load_model(model_path)
                logger.info(f"저장된 임베딩 모델 로드: {model_path}")
            else:
                # 새로운 모델은 문서 추가 시 학습
                logger.info("새로운 오프라인 임베딩 모델 준비")
                
            return True
            
        except Exception as e:
            logger.error(f"임베딩 모델 초기화 실패: {str(e)}")
            return False
            
    def process_docx(self, file_path: str) -> List[DocumentChunk]:
        """DOCX 파일 처리"""
        try:
            doc = Document(file_path)
            full_text = ""
            metadata = {
                "source": os.path.basename(file_path),
                "type": "docx",
                "processed_at": datetime.now().isoformat()
            }
            
            # 문서 내용 추출
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n\n"
                    
            # 테이블 내용 추출
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    full_text += "\n[테이블]\n" + table_text + "\n\n"
                    
            # 텍스트 분할
            chunks = self.text_splitter.split_text(full_text)
            
            # DocumentChunk 객체 생성
            doc_chunks = []
            for i, chunk_text in enumerate(chunks):
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata={**metadata, "chunk_index": i}
                )
                doc_chunks.append(chunk)
                
            return doc_chunks
            
        except Exception as e:
            logger.error(f"DOCX 처리 오류: {str(e)}")
            raise
            
    def _extract_table_text(self, table) -> str:
        """테이블 내용 추출"""
        table_text = ""
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            table_text += row_text + "\n"
        return table_text
        
    def create_embeddings(self, chunks: List[DocumentChunk]):
        """청크 임베딩 생성"""
        # 모델이 학습되지 않았다면 현재 문서로 학습
        if not self.embedding_model.trained:
            all_texts = [chunk.content for chunk in self.chunks_db] + [chunk.content for chunk in chunks]
            self.embedding_model.build_from_documents(all_texts)
            
        texts = [chunk.content for chunk in chunks]
        embeddings = self.embedding_model.encode(texts)
        
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding
            
        # 임베딩 인덱스 업데이트
        self._update_index(chunks)
        
        # TF-IDF 매트릭스 업데이트 (하이브리드 검색용)
        if self.config.use_hybrid_search:
            self._update_tfidf(chunks)
            
    def _update_index(self, new_chunks: List[DocumentChunk]):
        """FAISS 인덱스 업데이트"""
        self.chunks_db.extend(new_chunks)
        
        all_embeddings = np.array([chunk.embedding for chunk in self.chunks_db])
        dimension = all_embeddings.shape[1]
        
        # FAISS 인덱스 생성
        self.index = faiss.IndexFlatL2(dimension)
        self.index.add(all_embeddings.astype('float32'))
        
    def _update_tfidf(self, new_chunks: List[DocumentChunk]):
        """TF-IDF 매트릭스 업데이트"""
        all_texts = [chunk.content for chunk in self.chunks_db]
        
        self.tfidf_vectorizer = TfidfVectorizer(
            tokenizer=self._korean_tokenizer,
            max_features=5000,
            ngram_range=(1, 3)
        )
        self.tfidf_matrix = self.tfidf_vectorizer.fit_transform(all_texts)
        
    def _korean_tokenizer(self, text):
        """간단한 한글 토크나이저"""
        return self.embedding_model._tokenize_korean(text)
        
    def search(self, query: str, k: int = 5) -> List[Tuple[DocumentChunk, float]]:
        """하이브리드 검색 (벡터 + 키워드)"""
        if not self.index or not self.chunks_db:
            return []
            
        results = []
        
        # 1. 벡터 검색
        query_embedding = self.embedding_model.encode([query])[0]
        distances, indices = self.index.search(
            query_embedding.reshape(1, -1).astype('float32'), 
            min(k * 2, len(self.chunks_db))
        )
        
        vector_scores = {}
        for idx, dist in zip(indices[0], distances[0]):
            if idx < len(self.chunks_db):
                # 거리를 유사도로 변환
                similarity = 1 / (1 + dist)
                vector_scores[idx] = similarity
                
        # 2. 키워드 검색 (TF-IDF)
        keyword_scores = {}
        if self.config.use_hybrid_search and self.tfidf_matrix is not None:
            query_tfidf = self.tfidf_vectorizer.transform([query])
            keyword_similarities = cosine_similarity(query_tfidf, self.tfidf_matrix)[0]
            
            top_k_indices = np.argsort(keyword_similarities)[-k*2:][::-1]
            for idx in top_k_indices:
                if keyword_similarities[idx] > 0:
                    keyword_scores[idx] = keyword_similarities[idx]
                    
        # 3. 점수 결합
        all_indices = set(vector_scores.keys()) | set(keyword_scores.keys())
        combined_scores = {}
        
        for idx in all_indices:
            vector_score = vector_scores.get(idx, 0)
            keyword_score = keyword_scores.get(idx, 0)
            
            # 가중 평균
            combined_score = 0.6 * vector_score + 0.4 * keyword_score
            
            if combined_score >= self.config.min_relevance_score:
                combined_scores[idx] = combined_score
                
        # 4. 상위 k개 선택
        sorted_indices = sorted(combined_scores.items(), key=lambda x: x[1], reverse=True)[:k]
        
        for idx, score in sorted_indices:
            results.append((self.chunks_db[idx], score))
            
        return results
        
    def save_embedding_model(self, path: str):
        """임베딩 모델 저장"""
        if self.embedding_model.trained:
            self.embedding_model.save_model(path)

class KoreanRAGChat(ctk.CTk):
    """한글 기술문서 RAG 채팅 애플리케이션"""
    
    def __init__(self):
        super().__init__()
        
        # 기본 설정
        self.title("Korean Tech RAG Chat - Offline Edition")
        self.geometry("1400x900")
        
        # 테마 설정
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        # 컴포넌트 초기화
        self.model_manager = None
        self.rag_config = RAGConfig()
        self.doc_processor = DocumentProcessor(self.rag_config)
        self.messages = []
        self.current_streaming = None
        
        # UI 구성
        self.setup_ui()
        
        # 초기 메시지
        self.add_message("assistant", 
            "안녕하세요! 👋\n\n"
            "오프라인 한글 기술문서 RAG 시스템입니다.\n"
            "1. GGUF 모델을 로드하세요 (SOLAR-10.7B 추천)\n"
            "2. DOCX 문서를 업로드하세요\n"
            "3. 기술 문서에 대해 질문하세요!\n\n"
            "💡 폐쇄망 환경에서도 완벽하게 작동합니다!")
            
    def setup_ui(self):
        """UI 구성"""
        # 메인 그리드
        self.grid_columnconfigure(0, weight=0)  # 좌측 사이드바
        self.grid_columnconfigure(1, weight=1)  # 중앙 채팅
        self.grid_columnconfigure(2, weight=0)  # 우측 문서 패널
        self.grid_rowconfigure(0, weight=1)
        
        # 사이드바
        self.setup_sidebar()
        
        # 채팅 영역
        self.setup_chat_area()
        
        # 문서 패널
        self.setup_document_panel()
        
    def setup_sidebar(self):
        """사이드바 구성"""
        sidebar = ctk.CTkFrame(self, width=300, corner_radius=0)
        sidebar.grid(row=0, column=0, sticky="nsew")
        sidebar.grid_rowconfigure(10, weight=1)
        
        # 타이틀
        title = ctk.CTkLabel(
            sidebar,
            text="Korean Tech RAG",
            font=("Arial", 24, "bold")
        )
        title.grid(row=0, column=0, padx=20, pady=(20, 10))
        
        # 오프라인 모드 표시
        offline_label = ctk.CTkLabel(
            sidebar,
            text="🔒 오프라인 모드",
            font=("Arial", 14, "bold"),
            text_color="lightgreen"
        )
        offline_label.grid(row=1, column=0, padx=20, pady=(0, 10))
        
        # 모델 정보
        self.model_info = ctk.CTkLabel(
            sidebar,
            text="모델: 로드되지 않음",
            font=("Arial", 12),
            text_color="gray"
        )
        self.model_info.grid(row=2, column=0, padx=20, pady=(0, 20))
        
        # 모델 관련 버튼
        ctk.CTkButton(
            sidebar,
            text="GGUF 모델 로드",
            command=self.load_model,
            height=40,
            fg_color="green",
            hover_color="darkgreen"
        ).grid(row=3, column=0, padx=20, pady=5, sticky="ew")
        
        # RAG 설정 섹션
        ctk.CTkLabel(
            sidebar, 
            text="RAG 설정",
            font=("Arial", 16, "bold")
        ).grid(row=4, column=0, padx=20, pady=(20, 10))
        
        # 청크 크기
        ctk.CTkLabel(sidebar, text="청크 크기:").grid(row=5, column=0, padx=20, pady=5, sticky="w")
        self.chunk_size_var = ctk.IntVar(value=1000)
        ctk.CTkSlider(
            sidebar,
            from_=500,
            to=2000,
            number_of_steps=15,
            variable=self.chunk_size_var,
            command=self.update_chunk_size
        ).grid(row=6, column=0, padx=20, pady=5, sticky="ew")
        
        self.chunk_size_label = ctk.CTkLabel(sidebar, text="1000")
        self.chunk_size_label.grid(row=7, column=0, padx=20, pady=(0, 10))
        
        # 검색 결과 수
        ctk.CTkLabel(sidebar, text="검색 문서 수:").grid(row=8, column=0, padx=20, pady=5, sticky="w")
        self.max_chunks_var = ctk.IntVar(value=5)
        ctk.CTkSlider(
            sidebar,
            from_=1,
            to=10,
            number_of_steps=9,
            variable=self.max_chunks_var,
            command=self.update_max_chunks
        ).grid(row=9, column=0, padx=20, pady=5, sticky="ew")
        
        self.max_chunks_label = ctk.CTkLabel(sidebar, text="5")
        self.max_chunks_label.grid(row=10, column=0, padx=20, pady=(0, 20))
        
        # 하이브리드 검색 토글
        self.hybrid_search_var = ctk.BooleanVar(value=True)
        ctk.CTkCheckBox(
            sidebar,
            text="하이브리드 검색 사용",
            variable=self.hybrid_search_var,
            command=self.toggle_hybrid_search
        ).grid(row=11, column=0, padx=20, pady=10, sticky="w")
        
        # 임베딩 모델 저장/로드
        ctk.CTkButton(
            sidebar,
            text="임베딩 모델 저장",
            command=self.save_embedding_model,
            height=35
        ).grid(row=12, column=0, padx=20, pady=5, sticky="ew")
        
        ctk.CTkButton(
            sidebar,
            text="임베딩 모델 로드",
            command=self.load_embedding_model,
            height=35
        ).grid(row=13, column=0, padx=20, pady=5, sticky="ew")
        
        # 대화 초기화
        ctk.CTkButton(
            sidebar,
            text="대화 초기화",
            command=self.clear_chat,
            height=35,
            fg_color="red",
            hover_color="darkred"
        ).grid(row=14, column=0, padx=20, pady=(20, 20), sticky="ew")
        
    def setup_chat_area(self):
        """채팅 영역 구성"""
        chat_frame = ctk.CTkFrame(self)
        chat_frame.grid(row=0, column=1, sticky="nsew", padx=10, pady=10)
        chat_frame.grid_columnconfigure(0, weight=1)
        chat_frame.grid_rowconfigure(0, weight=1)
        
        # 채팅 스크롤 영역
        self.chat_scroll = ctk.CTkScrollableFrame(chat_frame)
        self.chat_scroll.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        # 입력 영역
        input_frame = ctk.CTkFrame(chat_frame)
        input_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 10))
        input_frame.grid_columnconfigure(0, weight=1)
        
        # 입력 텍스트박스
        self.input_text = ctk.CTkTextbox(
            input_frame,
            height=100,
            wrap="word",
            font=("Arial", 14)
        )
        self.input_text.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.input_text.bind("<Control-Return>", lambda e: self.send_message())
        
        # 버튼 프레임
        btn_frame = ctk.CTkFrame(input_frame)
        btn_frame.grid(row=0, column=1, sticky="ns")
        
        self.send_btn = ctk.CTkButton(
            btn_frame,
            text="전송",
            command=self.send_message,
            width=100,
            height=40,
            state="disabled"
        )
        self.send_btn.pack(pady=(0, 5))
        
        self.stop_btn = ctk.CTkButton(
            btn_frame,
            text="중지",
            command=self.stop_generation,
            width=100,
            height=40,
            fg_color="orange",
            hover_color="darkorange",
            state="disabled"
        )
        self.stop_btn.pack()
        
    def setup_document_panel(self):
        """문서 패널 구성"""
        doc_panel = ctk.CTkFrame(self, width=350, corner_radius=0)
        doc_panel.grid(row=0, column=2, sticky="nsew")
        doc_panel.grid_rowconfigure(2, weight=1)
        
        # 타이틀
        ctk.CTkLabel(
            doc_panel,
            text="문서 관리",
            font=("Arial", 20, "bold")
        ).grid(row=0, column=0, padx=20, pady=(20, 10))
        
        # 문서 업로드 버튼
        ctk.CTkButton(
            doc_panel,
            text="DOCX 문서 업로드",
            command=self.upload_documents,
            height=40,
            fg_color="blue",
            hover_color="darkblue"
        ).grid(row=1, column=0, padx=20, pady=10, sticky="ew")
        
        # 문서 목록
        self.doc_listbox = ctk.CTkTextbox(
            doc_panel,
            height=400,
            font=("Arial", 12)
        )
        self.doc_listbox.grid(row=2, column=0, padx=20, pady=10, sticky="nsew")
        
        # 문서 통계
        self.doc_stats = ctk.CTkLabel(
            doc_panel,
            text="문서: 0개\n청크: 0개\n임베딩: ❌",
            font=("Arial", 12),
            justify="left"
        )
        self.doc_stats.grid(row=3, column=0, padx=20, pady=10, sticky="w")
        
        # 문서 초기화 버튼
        ctk.CTkButton(
            doc_panel,
            text="문서 전체 삭제",
            command=self.clear_documents,
            height=35,
            fg_color="red",
            hover_color="darkred"
        ).grid(row=4, column=0, padx=20, pady=(10, 20), sticky="ew")
        
    def add_message(self, role: str, content: str, streaming=False, metadata=None):
        """메시지 추가"""
        # 메시지 컨테이너
        msg_frame = ctk.CTkFrame(self.chat_scroll, corner_radius=10)
        
        if role == "user":
            msg_frame.configure(fg_color=("gray85", "gray25"))
            msg_frame.pack(anchor="e", padx=(100, 10), pady=5, fill="x")
        else:
            msg_frame.configure(fg_color=("gray90", "gray20"))
            msg_frame.pack(anchor="w", padx=(10, 100), pady=5, fill="x")
        
        # 역할 라벨
        role_label = ctk.CTkLabel(
            msg_frame,
            text="You" if role == "user" else "AI",
            font=("Arial", 12, "bold"),
            text_color=("gray40", "gray60")
        )
        role_label.pack(anchor="w", padx=15, pady=(10, 0))
        
        # 메타데이터 표시 (참조 문서 등)
        if metadata and "sources" in metadata:
            sources_text = "📚 참조: " + ", ".join(metadata["sources"])
            sources_label = ctk.CTkLabel(
                msg_frame,
                text=sources_text,
                font=("Arial", 10),
                text_color=("blue", "lightblue")
            )
            sources_label.pack(anchor="w", padx=15, pady=(5, 0))
        
        # 메시지 라벨
        msg_label = ctk.CTkLabel(
            msg_frame,
            text=content,
            font=("Arial", 14),
            wraplength=700,
            justify="left"
        )
        msg_label.pack(anchor="w", padx=15, pady=(5, 10))
        
        # 메시지 저장
        if not streaming:
            self.messages.append({
                "role": role,
                "content": content,
                "timestamp": datetime.now(),
                "metadata": metadata
            })
        
        # 스크롤 다운
        self.chat_scroll._parent_canvas.yview_moveto(1.0)
        
        return msg_label
        
    def load_model(self):
        """GGUF 모델 로드"""
        filepath = filedialog.askopenfilename(
            title="GGUF 모델 선택 (SOLAR-10.7B 권장)",
            filetypes=[("GGUF Files", "*.gguf"), ("All Files", "*.*")]
        )
        
        if not filepath:
            return
            
        # 로딩 다이얼로그
        self.show_loading("모델 로드 중...")
        
        # 별도 스레드에서 로드
        threading.Thread(
            target=self._load_model_thread,
            args=(filepath,),
            daemon=True
        ).start()
        
    def _load_model_thread(self, filepath):
        """모델 로드 스레드"""
        try:
            # GPU 레이어 자동 감지
            n_gpu_layers = -1
            try:
                import torch
                if torch.cuda.is_available():
                    n_gpu_layers = 999
            except:
                n_gpu_layers = 0
            
            self.model_manager = Llama(
                model_path=filepath,
                n_ctx=32768,  # SOLAR는 32K 컨텍스트
                n_threads=8,
                n_gpu_layers=n_gpu_layers,
                verbose=False,
                use_mmap=True,
                use_mlock=False,
                n_batch=512,
                rope_scaling_type=1,
                mul_mat_q=True,
            )
            
            # UI 업데이트
            self.after(0, self._on_model_loaded, filepath)
            
        except Exception as e:
            self.after(0, self._on_model_error, str(e))
            
    def _on_model_loaded(self, filepath):
        """모델 로드 완료"""
        self.hide_loading()
        model_name = os.path.basename(filepath)
        self.model_info.configure(text=f"모델: {model_name}")
        self.send_btn.configure(state="normal")
        messagebox.showinfo("성공", "모델이 성공적으로 로드되었습니다!")
        
    def _on_model_error(self, error):
        """모델 로드 오류"""
        self.hide_loading()
        messagebox.showerror("오류", f"모델 로드 실패:\n{error}")
        
    def upload_documents(self):
        """문서 업로드"""
        filepaths = filedialog.askopenfilenames(
            title="DOCX 문서 선택 (다중 선택 가능)",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not filepaths:
            return
            
        # 로딩 표시
        self.show_loading(f"{len(filepaths)}개 문서 처리 중...")
        
        # 별도 스레드에서 처리
        threading.Thread(
            target=self._process_documents_thread,
            args=(filepaths,),
            daemon=True
        ).start()
        
    def _process_documents_thread(self, filepaths):
        """문서 처리 스레드"""
        try:
            processed_docs = []
            
            for filepath in filepaths:
                # DOCX 처리
                chunks = self.doc_processor.process_docx(filepath)
                processed_docs.append({
                    "path": filepath,
                    "name": os.path.basename(filepath),
                    "chunks": len(chunks)
                })
                
                # 임베딩 생성
                self.doc_processor.create_embeddings(chunks)
                
            # UI 업데이트
            self.after(0, self._on_documents_processed, processed_docs)
            
        except Exception as e:
            self.after(0, self._on_document_error, str(e))
            
    def _on_documents_processed(self, processed_docs):
        """문서 처리 완료"""
        self.hide_loading()
        
        # 문서 목록 업데이트
        current_text = self.doc_listbox.get("1.0", "end").strip()
        for doc in processed_docs:
            new_text = f"📄 {doc['name']} ({doc['chunks']} chunks)\n"
            if current_text:
                current_text += "\n" + new_text
            else:
                current_text = new_text
                
        self.doc_listbox.delete("1.0", "end")
        self.doc_listbox.insert("1.0", current_text)
        
        # 통계 업데이트
        total_chunks = len(self.doc_processor.chunks_db)
        embedding_status = "✅" if self.doc_processor.embedding_model.trained else "❌"
        self.doc_stats.configure(
            text=f"문서: {len(processed_docs)}개\n청크: {total_chunks}개\n임베딩: {embedding_status}"
        )
        
        messagebox.showinfo("성공", f"{len(processed_docs)}개 문서가 처리되었습니다!")
        
    def _on_document_error(self, error):
        """문서 처리 오류"""
        self.hide_loading()
        messagebox.showerror("오류", f"문서 처리 실패:\n{error}")
        
    def send_message(self):
        """메시지 전송"""
        content = self.input_text.get("1.0", "end").strip()
        if not content or not self.model_manager:
            return
            
        # 입력 초기화
        self.input_text.delete("1.0", "end")
        
        # 사용자 메시지 추가
        self.add_message("user", content)
        
        # UI 상태 변경
        self.send_btn.configure(state="disabled")
        self.stop_btn.configure(state="normal")
        
        # RAG 검색 및 응답 생성
        threading.Thread(
            target=self._generate_rag_response,
            args=(content,),
            daemon=True
        ).start()
        
    def _generate_rag_response(self, query: str):
        """RAG 응답 생성"""
        try:
            # 1. 관련 문서 검색
            relevant_chunks = []
            sources = []
            
            if self.doc_processor.chunks_db:
                search_results = self.doc_processor.search(
                    query, 
                    k=self.rag_config.max_chunks_per_query
                )
                
                for chunk, score in search_results:
                    relevant_chunks.append(f"[관련도: {score:.2f}]\n{chunk.content}")
                    source = chunk.metadata.get("source", "Unknown")
                    if source not in sources:
                        sources.append(source)
                        
            # 2. 프롬프트 생성
            context = "\n\n---\n\n".join(relevant_chunks) if relevant_chunks else "참조할 문서가 없습니다."
            
            # 대화 기록 포맷팅
            chat_history = ""
            for msg in self.messages[-6:]:  # 최근 6개 메시지
                if msg["role"] == "user":
                    chat_history += f"사용자: {msg['content']}\n"
                elif msg["role"] == "assistant":
                    chat_history += f"AI: {msg['content']}\n"
                    
            # 프롬프트 템플릿 선택
            template = RAGPromptTemplate.KOREAN_TECH_RAG
            if "쿼리" in query.lower() or "sql" in query.lower():
                template = RAGPromptTemplate.QUERY_FOCUSED_RAG
                
            prompt = template.value.format(
                system_prompt="한글 기술 문서를 정확하게 분석하고 답변합니다.",
                context=context,
                chat_history=chat_history,
                user_message=query
            )
            
            # 3. 스트리밍 응답 생성
            full_response = ""
            msg_label = None
            
            stream = self.model_manager(
                prompt,
                max_tokens=2048,
                temperature=0.3,  # 정확도 우선
                top_p=0.9,
                top_k=40,
                repeat_penalty=1.1,
                stream=True,
                stop=["사용자:", "Human:", "User:", "\n\n\n"]
            )
            
            for output in stream:
                token = output['choices'][0]['text']
                full_response += token
                
                # UI 업데이트
                if msg_label is None:
                    metadata = {"sources": sources} if sources else None
                    self.after(0, lambda: setattr(self, '_temp_label', 
                        self.add_message("assistant", token, streaming=True, metadata=metadata)))
                    time.sleep(0.1)
                    msg_label = getattr(self, '_temp_label', None)
                else:
                    self.after(0, lambda t=full_response: msg_label.configure(text=t))
                    
            # 최종 메시지 저장
            self.messages.append({
                "role": "assistant",
                "content": full_response,
                "timestamp": datetime.now(),
                "metadata": {"sources": sources} if sources else None
            })
            
        except Exception as e:
            error_msg = f"오류 발생: {str(e)}"
            self.after(0, lambda: self.add_message("assistant", error_msg))
            
        finally:
            # UI 상태 복원
            self.after(0, self._reset_ui_state)
            
    def stop_generation(self):
        """생성 중지"""
        # 구현 필요
        pass
        
    def _reset_ui_state(self):
        """UI 상태 초기화"""
        self.send_btn.configure(state="normal")
        self.stop_btn.configure(state="disabled")
        
    def clear_chat(self):
        """대화 초기화"""
        if messagebox.askyesno("확인", "대화를 초기화하시겠습니까?"):
            self.messages.clear()
            for widget in self.chat_scroll.winfo_children():
                widget.destroy()
            self.add_message("assistant", "대화가 초기화되었습니다.")
            
    def clear_documents(self):
        """문서 초기화"""
        if messagebox.askyesno("확인", "모든 문서를 삭제하시겠습니까?"):
            self.doc_processor.chunks_db.clear()
            self.doc_processor.index = None
            self.doc_processor.tfidf_matrix = None
            self.doc_processor.embedding_model = KoreanOfflineEmbedding(self.rag_config.embedding_dimension)
            self.doc_listbox.delete("1.0", "end")
            self.doc_stats.configure(text="문서: 0개\n청크: 0개\n임베딩: ❌")
            messagebox.showinfo("완료", "모든 문서가 삭제되었습니다.")
            
    def save_embedding_model(self):
        """임베딩 모델 저장"""
        if not self.doc_processor.embedding_model.trained:
            messagebox.showwarning("경고", "학습된 임베딩 모델이 없습니다.")
            return
            
        filepath = filedialog.asksaveasfilename(
            title="임베딩 모델 저장",
            defaultextension=".pkl",
            filetypes=[("Pickle Files", "*.pkl"), ("All Files", "*.*")]
        )
        
        if filepath:
            self.doc_processor.save_embedding_model(filepath)
            messagebox.showinfo("성공", "임베딩 모델이 저장되었습니다!")
            
    def load_embedding_model(self):
        """임베딩 모델 로드"""
        filepath = filedialog.askopenfilename(
            title="임베딩 모델 로드",
            filetypes=[("Pickle Files", "*.pkl"), ("All Files", "*.*")]
        )
        
        if filepath:
            success = self.doc_processor.initialize_embedding_model(filepath)
            if success:
                messagebox.showinfo("성공", "임베딩 모델이 로드되었습니다!")
                self.doc_stats.configure(
                    text=self.doc_stats.cget("text").replace("임베딩: ❌", "임베딩: ✅")
                )
            else:
                messagebox.showerror("오류", "임베딩 모델 로드 실패")
                
    def update_chunk_size(self, value):
        """청크 크기 업데이트"""
        self.chunk_size_label.configure(text=str(int(value)))
        self.rag_config.chunk_size = int(value)
        self.doc_processor.text_splitter.chunk_size = int(value)
        
    def update_max_chunks(self, value):
        """최대 청크 수 업데이트"""
        self.max_chunks_label.configure(text=str(int(value)))
        self.rag_config.max_chunks_per_query = int(value)
        
    def toggle_hybrid_search(self):
        """하이브리드 검색 토글"""
        self.rag_config.use_hybrid_search = self.hybrid_search_var.get()
        
    def show_loading(self, message):
        """로딩 표시"""
        self.loading_window = ctk.CTkToplevel(self)
        self.loading_window.title("처리 중")
        self.loading_window.geometry("300x150")
        self.loading_window.transient(self)
        self.loading_window.grab_set()
        
        # 중앙 정렬
        self.loading_window.update_idletasks()
        x = (self.loading_window.winfo_screenwidth() // 2) - 150
        y = (self.loading_window.winfo_screenheight() // 2) - 75
        self.loading_window.geometry(f"+{x}+{y}")
        
        ctk.CTkLabel(
            self.loading_window,
            text=message,
            font=("Arial", 16)
        ).pack(pady=30)
        
        progress = ctk.CTkProgressBar(self.loading_window, mode="indeterminate")
        progress.pack(padx=40)
        progress.start()
        
    def hide_loading(self):
        """로딩 숨기기"""
        if hasattr(self, 'loading_window'):
            self.loading_window.destroy()

def main():
    """메인 함수"""
    if not LLAMA_CPP_AVAILABLE:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror(
            "Error",
            "llama-cpp-python is not installed.\n\n"
            "Please install it using:\n"
            "pip install llama-cpp-python"
        )
        return
        
    if not RAG_LIBS_AVAILABLE:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror(
            "Error",
            "RAG libraries are not installed.\n\n"
            "Please install them using:\n"
            "pip install python-docx faiss-cpu scikit-learn numpy"
        )
        return
        
    # 필요한 디렉토리 생성
    for dir_name in ["models", "documents", "embeddings", "logs"]:
        Path(dir_name).mkdir(exist_ok=True)
        
    # 앱 실행
    app = KoreanRAGChat()
    app.mainloop()

if __name__ == "__main__":
    main()